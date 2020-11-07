from docx.api import Document
from utils.table_utils import table_to_df


class DocToExcel(object):
    def __init__(self, filename,):
        self.filename = filename
        self.document = Document(filename)
        self.outlist = []

    def __repr__(self):
        return f" The Document {self.filename} has {len(self.document.tables)} tables"

    def convert_table_to_df(self, table_list=""):
        if table_list == '':
            table_nos = list(range(len(self.document.tables)))
        else:
            table_nos = [int(i) for i in table_list.split()]
        print(f"{len(table_nos)} are to be extracted")
        for table in table_nos:
            input_table = self.document.tables[table]
            table_dataframe = table_to_df(input_table)
            self.outlist.append(table_dataframe)
        
        return self.outlist.copy()
