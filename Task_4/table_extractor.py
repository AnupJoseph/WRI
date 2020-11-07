import os
import pandas as pd
from docx.api import Document
import xlsxwriter
import fire


def table_to_df(table):
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)

    df = pd.DataFrame(data)
    print(df)
    return df


def save_to_excel(dataframe_list, outfolder='excels'):
    if not os.path.exists('excels'):
        os.makedirs('excels')
    writer = pd.ExcelWriter(
        'excels/Combined_table_list.xlsx', engine='xlsxwriter')
    for tables in dataframe_list:
        tables.to_excel(writer)
    print("Writing Done")
    writer.save()


def convert_table_to_df(document_name, table_nos=[]):
    document = Document(document_name)
    outlist = []
    if table_nos == '':
        table_nos = list(range(len(document.tables)))
    for table in table_nos:
        input_table = document.tables[table]

        table_dataframe = table_to_df(input_table)
        outlist.append(table_dataframe)
    return outlist


class Controller(object):
    def save_table(self, filename='/content/drive/My Drive/Tariff Order 2017.docx', table_nos=""):
        table_nos = str(table_nos)
        table_nos = list(table_nos.split())
        table_nos = [int(i) for i in table_nos]
        print(f"Tables to be converted {table_nos}")
        outlist = convert_table_to_df(filename, table_nos)
        save_to_excel(outlist)


if __name__ == "__main__":
    fire.Fire(Controller)
